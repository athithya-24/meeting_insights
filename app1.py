from flask import Flask, request, jsonify, render_template, send_file, Response
import os
import whisper
import dotenv
import google.generativeai as genai
import logging
from datetime import datetime, timedelta
import json
from werkzeug.utils import secure_filename
import tempfile
import shutil
import io
from collections import Counter
import re
import base64
import math

# Document generation imports
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors

try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not available. DOCX export will be disabled.")

# Visualization imports
try:
    import matplotlib
    matplotlib.use('Agg')  # Use non-interactive backend
    import matplotlib.pyplot as plt
    from wordcloud import WordCloud
    import numpy as np
    VISUALIZATION_AVAILABLE = True
except ImportError:
    VISUALIZATION_AVAILABLE = False
    logging.warning("Visualization libraries not available. Charts will be disabled.")

# Manually append the ffmpeg path so Python can access it
os.environ["PATH"] += os.pathsep + r"C:\Users\Athithya\Downloads\ffmpeg-7.1.1-essentials_build\ffmpeg-7.1.1-essentials_build\bin"

# Load .env file
dotenv.load_dotenv()

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['EXPORT_FOLDER'] = 'exports'
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB max file size
app.config['ALLOWED_EXTENSIONS'] = {'wav', 'mp3', 'webm', 'm4a', 'ogg', 'flac'}

# Create necessary directories
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['EXPORT_FOLDER'], exist_ok=True)
os.makedirs('meeting_history', exist_ok=True)

# Load Whisper model with error handling
try:
    whisper_model = whisper.load_model("base")
    logger.info("Whisper model loaded successfully")
except Exception as e:
    logger.error(f"Failed to load Whisper model: {e}")
    whisper_model = None

# Setup Gemini with error handling
try:
    genai.configure(api_key=os.getenv("GEMINI_API_KEY"))
    gemini_model = genai.GenerativeModel("gemini-1.5-flash")
    logger.info("Gemini model configured successfully")
except Exception as e:
    logger.error(f"Failed to configure Gemini: {e}")
    gemini_model = None

def allowed_file(filename):
    """Check if the uploaded file has an allowed extension"""
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

def save_meeting_history(transcript, summary, metadata=None, analysis_data=None):
    """Save meeting data to history for future reference"""
    try:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"meeting_{timestamp}.json"
        filepath = os.path.join('meeting_history', filename)
        
        meeting_data = {
            'timestamp': datetime.now().isoformat(),
            'transcript': transcript,
            'summary': summary,
            'metadata': metadata or {},
            'analysis_data': analysis_data or {},
            'word_count': len(transcript.split()),
            'duration': metadata.get('duration', 0) if metadata else 0
        }
        
        with open(filepath, 'w', encoding='utf-8') as f:
            json.dump(meeting_data, f, indent=2, ensure_ascii=False)
        
        logger.info(f"Meeting history saved: {filename}")
        return filename
    except Exception as e:
        logger.error(f"Failed to save meeting history: {e}")
        return None

def generate_word_frequency(transcript, top_n=20):
    """Generate word frequency analysis"""
    try:
        # Clean and tokenize text
        text = re.sub(r'[^\w\s]', '', transcript.lower())
        words = text.split()
        
        # Remove common stop words
        stop_words = {'the', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by', 'a', 'an', 'is', 'are', 'was', 'were', 'be', 'been', 'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would', 'should', 'could', 'can', 'may', 'might', 'must', 'shall', 'this', 'that', 'these', 'those', 'i', 'you', 'he', 'she', 'it', 'we', 'they', 'me', 'him', 'her', 'us', 'them', 'my', 'your', 'his', 'her', 'its', 'our', 'their'}
        
        filtered_words = [word for word in words if len(word) > 2 and word not in stop_words]
        
        # Count frequencies
        word_freq = Counter(filtered_words)
        top_words = dict(word_freq.most_common(top_n))
        
        return top_words
    except Exception as e:
        logger.error(f"Error generating word frequency: {e}")
        return {}

def generate_word_cloud(transcript):
    """Generate word cloud image"""
    if not VISUALIZATION_AVAILABLE:
        return None
    
    try:
        # Generate word frequency
        word_freq = generate_word_frequency(transcript, 100)
        
        if not word_freq:
            return None
        
        # Create word cloud
        wordcloud = WordCloud(
            width=800, 
            height=400, 
            background_color='white',
            colormap='viridis',
            max_words=100
        ).generate_from_frequencies(word_freq)
        
        # Save to bytes
        img_buffer = io.BytesIO()
        plt.figure(figsize=(10, 5))
        plt.imshow(wordcloud, interpolation='bilinear')
        plt.axis('off')
        plt.tight_layout(pad=0)
        plt.savefig(img_buffer, format='png', bbox_inches='tight', dpi=150)
        plt.close()
        
        img_buffer.seek(0)
        img_base64 = base64.b64encode(img_buffer.getvalue()).decode()
        
        return f"data:image/png;base64,{img_base64}"
    except Exception as e:
        logger.error(f"Error generating word cloud: {e}")
        return None

def generate_sentiment_chart(sentiment_data):
    """Generate sentiment analysis chart"""
    if not VISUALIZATION_AVAILABLE or not sentiment_data:
        return None
    
    try:
        segments = sentiment_data.get('segments', [])
        if not segments:
            return None
        
        times = [seg['time'] for seg in segments]
        sentiments = [seg['sentiment_score'] for seg in segments]
        colors_map = ['red' if s < -0.1 else 'green' if s > 0.1 else 'gray' for s in sentiments]
        
        plt.figure(figsize=(12, 6))
        plt.bar(range(len(times)), sentiments, color=colors_map, alpha=0.7)
        plt.axhline(y=0, color='black', linestyle='-', alpha=0.3)
        plt.xlabel('Time Segments')
        plt.ylabel('Sentiment Score')
        plt.title('Sentiment Analysis Over Time')
        plt.xticks(range(len(times)), [f"{int(t//60)}:{int(t%60):02d}" for t in times], rotation=45)
        
        # Add legend
        plt.axhline(y=0, color='red', alpha=0.7, label='Negative')
        plt.axhline(y=0, color='gray', alpha=0.7, label='Neutral')
        plt.axhline(y=0, color='green', alpha=0.7, label='Positive')
        plt.legend()
        
        img_buffer = io.BytesIO()
        plt.tight_layout()
        plt.savefig(img_buffer, format='png', bbox_inches='tight', dpi=150)
        plt.close()
        
        img_buffer.seek(0)
        img_base64 = base64.b64encode(img_buffer.getvalue()).decode()
        
        return f"data:image/png;base64,{img_base64}"
    except Exception as e:
        logger.error(f"Error generating sentiment chart: {e}")
        return None

def analyze_sentiment_by_segments(transcript_with_timestamps):
    """Analyze sentiment for different segments of the transcript"""
    try:
        if not gemini_model:
            return {'segments': [], 'overall': 'neutral'}
        
        # Split transcript into segments (every 30 seconds or by sentences)
        segments = []
        if isinstance(transcript_with_timestamps, dict) and 'segments' in transcript_with_timestamps:
            # Use Whisper segments if available
            whisper_segments = transcript_with_timestamps['segments']
            for i, segment in enumerate(whisper_segments):
                segments.append({
                    'text': segment['text'],
                    'start': segment['start'],
                    'end': segment['end'],
                    'index': i
                })
        else:
            # Fallback: split by sentences
            sentences = re.split(r'[.!?]+', transcript_with_timestamps)
            for i, sentence in enumerate(sentences):
                if sentence.strip():
                    segments.append({
                        'text': sentence.strip(),
                        'start': i * 10,  # Estimate 10 seconds per sentence
                        'end': (i + 1) * 10,
                        'index': i
                    })
        
        # Analyze sentiment for each segment
        segment_analysis = []
        for segment in segments[:20]:  # Limit to first 20 segments to avoid API limits
            try:
                prompt = f"""
                Analyze the sentiment of this text segment on a scale from -1 (very negative) to +1 (very positive):
                
                Text: "{segment['text']}"
                
                Respond with only a number between -1 and 1, followed by a brief reason (max 10 words).
                Format: "0.3 - Optimistic discussion about future plans"
                """
                
                response = gemini_model.generate_content(prompt)
                response_text = response.text.strip()
                
                # Parse response
                parts = response_text.split(' - ', 1)
                score = float(parts[0]) if parts[0].replace('-', '').replace('.', '').isdigit() else 0
                reason = parts[1] if len(parts) > 1 else 'No specific reason'
                
                segment_analysis.append({
                    'time': segment['start'],
                    'text': segment['text'][:100] + '...' if len(segment['text']) > 100 else segment['text'],
                    'sentiment_score': max(-1, min(1, score)),  # Clamp between -1 and 1
                    'reason': reason
                })
            except Exception as e:
                logger.warning(f"Failed to analyze segment sentiment: {e}")
                segment_analysis.append({
                    'time': segment['start'],
                    'text': segment['text'][:100] + '...' if len(segment['text']) > 100 else segment['text'],
                    'sentiment_score': 0,
                    'reason': 'Analysis failed'
                })
        
        # Calculate overall sentiment
        if segment_analysis:
            overall_score = sum(seg['sentiment_score'] for seg in segment_analysis) / len(segment_analysis)
            if overall_score > 0.1:
                overall = 'positive'
            elif overall_score < -0.1:
                overall = 'negative'
            else:
                overall = 'neutral'
        else:
            overall = 'neutral'
        
        return {
            'segments': segment_analysis,
            'overall': overall,
            'average_score': overall_score if segment_analysis else 0
        }
    except Exception as e:
        logger.error(f"Error in sentiment analysis: {e}")
        return {'segments': [], 'overall': 'neutral', 'average_score': 0}

def generate_time_based_summary(transcript_result, interval_minutes=2):
    """Generate summaries for time-based segments"""
    try:
        if not gemini_model:
            return []
        
        summaries = []
        
        # Use Whisper segments if available
        if 'segments' in transcript_result:
            segments = transcript_result['segments']
            interval_seconds = interval_minutes * 60
            
            current_time = 0
            current_text = ""
            segment_index = 1
            
            for segment in segments:
                segment_start = segment['start']
                segment_text = segment['text']
                
                # If we've reached the interval, summarize current text
                if segment_start >= current_time + interval_seconds:
                    if current_text.strip():
                        try:
                            prompt = f"""
                            Provide a brief summary (max 2-3 sentences) of this {interval_minutes}-minute segment from a meeting:
                            
                            Time: {int(current_time//60)}:{int(current_time%60):02d} - {int((current_time + interval_seconds)//60)}:{int((current_time + interval_seconds)%60):02d}
                            
                            Content: {current_text}
                            """
                            
                            response = gemini_model.generate_content(prompt)
                            summary = response.text.strip()
                            
                            summaries.append({
                                'segment': segment_index,
                                'start_time': current_time,
                                'end_time': current_time + interval_seconds,
                                'duration_minutes': interval_minutes,
                                'summary': summary,
                                'word_count': len(current_text.split())
                            })
                            
                            segment_index += 1
                        except Exception as e:
                            logger.warning(f"Failed to generate segment summary: {e}")
                    
                    current_time += interval_seconds
                    current_text = segment_text
                else:
                    current_text += " " + segment_text
            
            # Handle remaining text
            if current_text.strip():
                try:
                    prompt = f"""
                    Provide a brief summary (max 2-3 sentences) of this final segment from a meeting:
                    
                    Time: {int(current_time//60)}:{int(current_time%60):02d} - End
                    
                    Content: {current_text}
                    """
                    
                    response = gemini_model.generate_content(prompt)
                    summary = response.text.strip()
                    
                    summaries.append({
                        'segment': segment_index,
                        'start_time': current_time,
                        'end_time': transcript_result.get('duration', current_time + 60),
                        'duration_minutes': interval_minutes,
                        'summary': summary,
                        'word_count': len(current_text.split())
                    })
                except Exception as e:
                    logger.warning(f"Failed to generate final segment summary: {e}")
        
        return summaries
    except Exception as e:
        logger.error(f"Error generating time-based summaries: {e}")
        return []

def generate_srt_subtitles(transcript_result):
    """Generate SRT subtitle file content"""
    try:
        if 'segments' not in transcript_result:
            return None
        
        srt_content = ""
        for i, segment in enumerate(transcript_result['segments'], 1):
            start_time = segment['start']
            end_time = segment['end']
            text = segment['text'].strip()
            
            # Convert seconds to SRT time format (HH:MM:SS,mmm)
            def seconds_to_srt_time(seconds):
                hours = int(seconds // 3600)
                minutes = int((seconds % 3600) // 60)
                secs = int(seconds % 60)
                millis = int((seconds - int(seconds)) * 1000)
                return f"{hours:02d}:{minutes:02d}:{secs:02d},{millis:03d}"
            
            start_srt = seconds_to_srt_time(start_time)
            end_srt = seconds_to_srt_time(end_time)
            
            srt_content += f"{i}\n{start_srt} --> {end_srt}\n{text}\n\n"
        
        return srt_content
    except Exception as e:
        logger.error(f"Error generating SRT subtitles: {e}")
        return None

def export_to_pdf(content, title="Meeting Report"):
    """Export content to PDF"""
    try:
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            alignment=1  # Center alignment
        )
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 20))
        
        # Content
        content_style = ParagraphStyle(
            'CustomContent',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=12,
            leftIndent=0,
            rightIndent=0
        )
        
        # Split content into paragraphs
        paragraphs = content.split('\n\n')
        for para in paragraphs:
            if para.strip():
                story.append(Paragraph(para.strip(), content_style))
                story.append(Spacer(1, 12))
        
        doc.build(story)
        buffer.seek(0)
        return buffer
    except Exception as e:
        logger.error(f"Error generating PDF: {e}")
        return None

def export_to_docx(content, title="Meeting Report"):
    """Export content to DOCX"""
    if not DOCX_AVAILABLE:
        return None
    
    try:
        doc = Document()
        
        # Add title
        title_para = doc.add_heading(title, 0)
        title_para.alignment = 1  # Center alignment
        
        # Add content
        paragraphs = content.split('\n\n')
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para.strip())
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    except Exception as e:
        logger.error(f"Error generating DOCX: {e}")
        return None

def generate_enhanced_summary(transcript, transcript_result=None):
    """Generate enhanced summary with multiple perspectives"""
    try:
        prompts = {
            'summary': f"""
            Please provide a comprehensive summary of this meeting transcript:
            
            {transcript}
            
            Include:
            - Main topics discussed
            - Key decisions made
            - Action items (if any)
            - Important points raised
            """,
            
            'action_items': f"""
            From this meeting transcript, extract any action items, tasks, or follow-ups mentioned:
            
            {transcript}
            
            Format as a bulleted list. If no action items are found, respond with "No specific action items identified."
            """,
            
            'sentiment': f"""
            Analyze the overall tone and sentiment of this meeting:
            
            {transcript}
            
            Provide a brief assessment of:
            - Overall mood (positive, neutral, negative)
            - Level of engagement
            - Any concerns or conflicts mentioned
            """
        }
        
        results = {}
        for key, prompt in prompts.items():
            try:
                response = gemini_model.generate_content(prompt)
                results[key] = response.text
            except Exception as e:
                logger.error(f"Failed to generate {key}: {e}")
                results[key] = f"Error generating {key}: {str(e)}"
        
        # Generate additional analysis
        word_frequency = generate_word_frequency(transcript)
        word_cloud_img = generate_word_cloud(transcript)
        
        # Detailed sentiment analysis
        sentiment_analysis = analyze_sentiment_by_segments(transcript_result or transcript)
        sentiment_chart = generate_sentiment_chart(sentiment_analysis)
        
        # Time-based summaries
        time_summaries = generate_time_based_summary(transcript_result or {'segments': []})
        
        results.update({
            'word_frequency': word_frequency,
            'word_cloud': word_cloud_img,
            'detailed_sentiment': sentiment_analysis,
            'sentiment_chart': sentiment_chart,
            'time_based_summaries': time_summaries
        })
        
        return results
    except Exception as e:
        logger.error(f"Failed to generate enhanced summary: {e}")
        return {
            'summary': f"Error generating summary: {str(e)}",
            'action_items': "Error extracting action items",
            'sentiment': "Error analyzing sentiment",
            'word_frequency': {},
            'word_cloud': None,
            'detailed_sentiment': {'segments': [], 'overall': 'neutral'},
            'sentiment_chart': None,
            'time_based_summaries': []
        }

@app.route('/')
def index():
    return render_template("index2.html")

@app.route('/health')
def health_check():
    """Health check endpoint"""
    status = {
        'status': 'healthy',
        'whisper_model': whisper_model is not None,
        'gemini_model': gemini_model is not None,
        'visualization_available': VISUALIZATION_AVAILABLE,
        'docx_available': DOCX_AVAILABLE,
        'timestamp': datetime.now().isoformat()
    }
    return jsonify(status)

@app.route('/process', methods=['POST'])
def process_audio():
    """Process uploaded audio file"""
    try:
        # Validation checks
        if whisper_model is None:
            return jsonify({'error': 'Whisper model not available'}), 500
        
        if gemini_model is None:
            return jsonify({'error': 'Gemini model not available'}), 500
        
        if 'audio_data' not in request.files:
            return jsonify({'error': 'No audio file provided'}), 400
        
        audio = request.files['audio_data']
        
        if audio.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        # Create temporary file with proper extension
        with tempfile.NamedTemporaryFile(delete=False, suffix='.webm') as temp_file:
            audio_path = temp_file.name
            audio.save(audio_path)
        
        try:
            logger.info(f"Processing audio file: {audio_path}")
            
            # Transcription with error handling and detailed results
            result = whisper_model.transcribe(audio_path, word_timestamps=True)
            transcript = result["text"].strip()
            
            if not transcript:
                return jsonify({'error': 'No speech detected in audio'}), 400
            
            # Get audio duration and other metadata
            metadata = {
                'duration': result.get('duration', 0),
                'language': result.get('language', 'unknown'),
                'processing_time': datetime.now().isoformat()
            }
            
            # Enhanced summarization with full analysis
            summary_data = generate_enhanced_summary(transcript, result)
            
            # Generate SRT subtitles
            srt_content = generate_srt_subtitles(result)
            
            # Save to history with analysis data
            history_file = save_meeting_history(
                transcript, 
                summary_data['summary'], 
                metadata, 
                {
                    'word_frequency': summary_data.get('word_frequency', {}),
                    'sentiment_analysis': summary_data.get('detailed_sentiment', {}),
                    'time_summaries': summary_data.get('time_based_summaries', []),
                    'srt_subtitles': srt_content
                }
            )
            
            # Prepare response
            response_data = {
                'transcript': transcript,
                'summary': summary_data['summary'],
                'action_items': summary_data.get('action_items', ''),
                'sentiment_analysis': summary_data.get('sentiment', ''),
                'detailed_sentiment': summary_data.get('detailed_sentiment', {}),
                'word_frequency': summary_data.get('word_frequency', {}),
                'word_cloud': summary_data.get('word_cloud'),
                'sentiment_chart': summary_data.get('sentiment_chart'),
                'time_based_summaries': summary_data.get('time_based_summaries', []),
                'srt_subtitles': srt_content,
                'metadata': metadata,
                'history_file': history_file,
                'word_count': len(transcript.split()),
                'processing_status': 'success'
            }
            
            logger.info("Audio processing completed successfully")
            return jsonify(response_data)
            
        finally:
            # Clean up temporary file
            try:
                os.unlink(audio_path)
            except Exception as e:
                logger.warning(f"Failed to delete temporary file: {e}")
    
    except Exception as e:
        logger.error(f"Error processing audio: {e}")
        return jsonify({
            'error': 'Internal server error during processing',
            'details': str(e)
        }), 500

@app.route('/export/<format>/<content_type>')
def export_content(format, content_type):
    """Export content in various formats"""
    try:
        # Get content from request parameters
        content = request.args.get('content', '')
        title = request.args.get('title', 'Meeting Export')
        
        if not content:
            return jsonify({'error': 'No content provided'}), 400
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        if format.lower() == 'pdf':
            buffer = export_to_pdf(content, title)
            if buffer:
                filename = f"{content_type}_{timestamp}.pdf"
                return send_file(
                    buffer,
                    as_attachment=True,
                    download_name=filename,
                    mimetype='application/pdf'
                )
        
        elif format.lower() == 'docx':
            if not DOCX_AVAILABLE:
                return jsonify({'error': 'DOCX export not available'}), 400
            
            buffer = export_to_docx(content, title)
            if buffer:
                filename = f"{content_type}_{timestamp}.docx"
                return send_file(
                    buffer,
                    as_attachment=True,
                    download_name=filename,
                    mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )
        
        elif format.lower() == 'txt':
            filename = f"{content_type}_{timestamp}.txt"
            return Response(
                content,
                mimetype='text/plain',
                headers={'Content-Disposition': f'attachment; filename={filename}'}
            )
        
        elif format.lower() == 'srt':
            filename = f"subtitles_{timestamp}.srt"
            return Response(
                content,
                mimetype='text/plain',
                headers={'Content-Disposition': f'attachment; filename={filename}'}
            )
        
        else:
            return jsonify({'error': 'Unsupported export format'}), 400
        
        return jsonify({'error': 'Export failed'}), 500
    
    except Exception as e:
        logger.error(f"Error exporting content: {e}")
        return jsonify({'error': 'Export failed', 'details': str(e)}), 500

@app.route('/history')
def get_meeting_history():
    """Get list of saved meetings"""
    try:
        history_files = []
        history_dir = 'meeting_history'
        
        if os.path.exists(history_dir):
            for filename in os.listdir(history_dir):
                if filename.endswith('.json'):
                    filepath = os.path.join(history_dir, filename)
                    try:
                        with open(filepath, 'r', encoding='utf-8') as f:
                            data = json.load(f)
                            history_files.append({
                                'filename': filename,
                                'timestamp': data.get('timestamp', ''),
                                'word_count': data.get('word_count', 0),
                                'duration': data.get('duration', 0)
                            })
                    except Exception as e:
                        logger.warning(f"Failed to read history file {filename}: {e}")
        
        # Sort by timestamp (newest first)
        history_files.sort(key=lambda x: x['timestamp'], reverse=True)
        
        return jsonify({
            'history': history_files,
            'total_meetings': len(history_files)
        })
    
    except Exception as e:
        logger.error(f"Error retrieving meeting history: {e}")
        return jsonify({'error': 'Failed to retrieve meeting history'}), 500

@app.route('/history/<filename>')
def get_meeting_details(filename):
    """Get details of a specific meeting"""
    try:
        # Sanitize filename
        safe_filename = secure_filename(filename)
        filepath = os.path.join('meeting_history', safe_filename)
        
        if not os.path.exists(filepath):
            return jsonify({'error': 'Meeting not found'}), 404
        
        with open(filepath, 'r', encoding='utf-8') as f:
            meeting_data = json.load(f)
        
        return jsonify(meeting_data)
    
    except Exception as e:
        logger.error(f"Error retrieving meeting details: {e}")
        return jsonify({'error': 'Failed to retrieve meeting details'}), 500

@app.errorhandler(413)
def too_large(e):
    return jsonify({'error': 'File too large. Maximum size is 50MB.'}), 413

@app.errorhandler(500)
def internal_error(e):
    logger.error(f"Internal server error: {e}")
    return jsonify({'error': 'Internal server error'}), 500

if __name__ == '__main__':
    logger.info("Starting Meeting Insights Pro server...")
    logger.info(f"Whisper model status: {'Loaded' if whisper_model else 'Failed'}")
    logger.info(f"Gemini model status: {'Configured' if gemini_model else 'Failed'}")
    logger.info(f"Visualization available: {VISUALIZATION_AVAILABLE}")
    logger.info(f"DOCX export available: {DOCX_AVAILABLE}")
    app.run(host='0.0.0.0', port=5000, debug=True)